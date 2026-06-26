# -*- coding: utf-8 -*-
"""移住支援関連（第2弾）docxテンプレートを生成するスクリプト。
リポジトリルートで `python tools/create_ijuu_templates2.py` を実行する。
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.join(os.path.dirname(__file__), "..")


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_table_border(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def label_cell(cell, text):
    set_cell_bg(cell, 'EBF0F8')
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.bold = True
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def value_cell(cell, placeholder):
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(placeholder)
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)


def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        cells = row.cells
        if col_idx < len(cells):
            cells[col_idx].width = Cm(width_cm)


def page_setup(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)


# ──────────────────────────────────────────────────────────────
# 1. 移住者向け住宅購入支援事業補助金交付申請書
# ──────────────────────────────────────────────────────────────
def create_jutaku_konyuu():
    doc = Document()
    page_setup(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.add_run('様式第１号（第４条関係）').font.size = Pt(9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(2)
    p.add_run('{{ shinsei_date }}').font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run('下関市移住者向け住宅購入支援事業補助金　交付申請書')
    run.font.size = Pt(15)
    run.font.bold = True

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.add_run('下関市長　様').font.size = Pt(11)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.add_run(
        '　下関市移住者向け住宅購入支援事業補助金交付要綱第４条の規定により、'
        '下記のとおり補助金の交付を申請します。'
    ).font.size = Pt(10)

    # 申請者情報
    add_section_heading(doc, '■ 申請者情報')
    tbl = doc.add_table(rows=4, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(tbl)
    data = [
        ('氏　　名', '{{ shimei }}', 'ふりがな', '{{ shimei_kana }}'),
        ('現住所', '{{ genjusho }}', '電話番号', '{{ denwa }}'),
        ('転入年月日', '{{ tennyu_date }}', '転入前住所', '{{ mae_jusho }}'),
        ('市外在住期間', '{{ zaiju_kikan }}', '申請区分', '{{ shinsei_kubun }}'),
    ]
    for i, (l1, v1, l2, v2) in enumerate(data):
        row = tbl.rows[i]
        label_cell(row.cells[0], l1)
        value_cell(row.cells[1], v1)
        label_cell(row.cells[2], l2)
        value_cell(row.cells[3], v2)
    set_col_width(tbl, 0, 3.0); set_col_width(tbl, 1, 5.5)
    set_col_width(tbl, 2, 3.0); set_col_width(tbl, 3, 5.5)

    doc.add_paragraph()

    # 住宅情報
    add_section_heading(doc, '■ 住宅情報')
    tbl2 = doc.add_table(rows=5, cols=4)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(tbl2)
    data2 = [
        ('住宅の種別', '{{ jutaku_shubetsu }}', '住宅の所在地', '{{ jutaku_shozaichi }}'),
        ('建築・購入費（円）', '{{ kenchiku_hi }}', '補助申請額（円）', '{{ hoshu_gaku }}'),
        ('契約締結日', '{{ keiyaku_date }}', '完成（引渡し）予定日', '{{ kansei_date }}'),
        ('中学生以下の子の数', '{{ kodomo_count }}人', '三世代同居・近居', '{{ sansetai }}'),
        ('居住誘導区域内か', '{{ yudo_kuiki }}', '', ''),
    ]
    for i, (l1, v1, l2, v2) in enumerate(data2):
        row = tbl2.rows[i]
        label_cell(row.cells[0], l1)
        value_cell(row.cells[1], v1)
        label_cell(row.cells[2], l2)
        value_cell(row.cells[3], v2)
    set_col_width(tbl2, 0, 3.0); set_col_width(tbl2, 1, 5.5)
    set_col_width(tbl2, 2, 3.0); set_col_width(tbl2, 3, 5.5)

    path = os.path.join(ROOT, 'jutaku_konyuu_template.docx')
    doc.save(path)
    print('ok: jutaku_konyuu_template.docx')


# ──────────────────────────────────────────────────────────────
# 2. やまぐち創生テレワーク移住支援事業補助金交付申請書
# ──────────────────────────────────────────────────────────────
def create_telework_ijuu():
    doc = Document()
    page_setup(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.add_run('下関市様式第１号（第４条関係）').font.size = Pt(9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(2)
    p.add_run('{{ shinsei_date }}').font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run('やまぐち創生テレワーク移住支援事業補助金　交付申請書')
    run.font.size = Pt(14)
    run.font.bold = True

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.add_run('下関市長　様').font.size = Pt(11)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.add_run(
        '　下関市やまぐち創生テレワーク移住支援事業補助金交付要綱第４条の規定により、'
        '下記のとおり補助金の交付を申請します。'
    ).font.size = Pt(10)

    # 申請者情報
    add_section_heading(doc, '■ 申請者情報')
    tbl = doc.add_table(rows=5, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(tbl)
    data = [
        ('氏　　名', '{{ shimei }}', 'ふりがな', '{{ shimei_kana }}'),
        ('現住所（下関市内）', '{{ jusho }}', '電話番号', '{{ denwa }}'),
        ('生年月日', '{{ seinengappi }}', '年　　齢', '{{ nenrei }}歳'),
        ('転入年月日', '{{ tennyu_date }}', '転入前住所\n（大都市圏の住所）', '{{ mae_jusho }}'),
        ('大都市圏での在住期間', '{{ zaiju_kikan }}', '転入前の職業', '{{ mae_shokugyo }}'),
    ]
    for i, (l1, v1, l2, v2) in enumerate(data):
        row = tbl.rows[i]
        label_cell(row.cells[0], l1)
        value_cell(row.cells[1], v1)
        label_cell(row.cells[2], l2)
        value_cell(row.cells[3], v2)
    set_col_width(tbl, 0, 3.0); set_col_width(tbl, 1, 5.5)
    set_col_width(tbl, 2, 3.0); set_col_width(tbl, 3, 5.5)

    doc.add_paragraph()

    # テレワーク就業情報
    add_section_heading(doc, '■ テレワーク就業情報')
    tbl2 = doc.add_table(rows=5, cols=4)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(tbl2)
    data2 = [
        ('勤務先（移住前の勤務先）', '{{ kinmusaki_name }}', '勤務先所在地', '{{ kinmusaki_jusho }}'),
        ('テレワーク開始日', '{{ telework_start_date }}', '週あたり\nテレワーク時間数', '{{ telework_jikan }}時間以上'),
        ('雇用形態', '{{ koyo_keitai }}', '', ''),
        ('テレワーク業務内容', '{{ telework_naiyou }}', '', ''),
        ('帯同する18歳未満の子の数', '{{ kodomo_count }}人', '', ''),
    ]
    for i, (l1, v1, l2, v2) in enumerate(data2):
        row = tbl2.rows[i]
        label_cell(row.cells[0], l1)
        value_cell(row.cells[1], v1)
        label_cell(row.cells[2], l2)
        value_cell(row.cells[3], v2)
    set_col_width(tbl2, 0, 3.0); set_col_width(tbl2, 1, 5.5)
    set_col_width(tbl2, 2, 3.0); set_col_width(tbl2, 3, 5.5)

    path = os.path.join(ROOT, 'telework_ijuu_template.docx')
    doc.save(path)
    print('ok: telework_ijuu_template.docx')


# ──────────────────────────────────────────────────────────────
# 3. YY!ターン支援交通費補助金申請書
# ──────────────────────────────────────────────────────────────
def create_yytan_kotsuhi():
    doc = Document()
    page_setup(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.add_run('別記第１号様式（第６条関係）').font.size = Pt(9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(2)
    p.add_run('{{ shinsei_date }}').font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run('YY!ターン支援交通費補助金　申請書')
    run.font.size = Pt(16)
    run.font.bold = True

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.add_run('下関市長　様').font.size = Pt(11)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.add_run(
        '　YY!ターン支援交通費補助金交付要綱第６条の規定により、'
        '下記のとおり補助金の交付を申請します。'
    ).font.size = Pt(10)

    # 申請者情報
    add_section_heading(doc, '■ 申請者情報')
    tbl = doc.add_table(rows=4, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(tbl)
    data = [
        ('氏　　名', '{{ shimei }}', 'ふりがな', '{{ shimei_kana }}'),
        ('現住所（山口県外）', '{{ jusho }}', '電話番号', '{{ denwa }}'),
        ('生年月日', '{{ seinengappi }}', '年　　齢', '{{ nenrei }}歳'),
        ('職　　業', '{{ shokugyo }}', '', ''),
    ]
    for i, (l1, v1, l2, v2) in enumerate(data):
        row = tbl.rows[i]
        label_cell(row.cells[0], l1)
        value_cell(row.cells[1], v1)
        label_cell(row.cells[2], l2)
        value_cell(row.cells[3], v2)
    set_col_width(tbl, 0, 3.0); set_col_width(tbl, 1, 5.5)
    set_col_width(tbl, 2, 3.0); set_col_width(tbl, 3, 5.5)

    doc.add_paragraph()

    # 移住活動情報
    add_section_heading(doc, '■ 移住活動情報')
    tbl2 = doc.add_table(rows=4, cols=4)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(tbl2)
    data2 = [
        ('活動実施日', '{{ katsudo_date }}', '活動場所（訪問先）', '{{ katsudo_basho }}'),
        ('活動内容', '{{ katsudo_naiyou }}', '', ''),
        ('往路交通手段\n（居住地→下関）', '{{ ofu_shudan }}', '往路交通費（円）', '{{ ofu_hi }}'),
        ('復路交通手段\n（下関→居住地）', '{{ fukuro_shudan }}', '復路交通費（円）', '{{ fukuro_hi }}'),
    ]
    for i, (l1, v1, l2, v2) in enumerate(data2):
        row = tbl2.rows[i]
        label_cell(row.cells[0], l1)
        value_cell(row.cells[1], v1)
        label_cell(row.cells[2], l2)
        value_cell(row.cells[3], v2)
    set_col_width(tbl2, 0, 3.0); set_col_width(tbl2, 1, 5.5)
    set_col_width(tbl2, 2, 3.0); set_col_width(tbl2, 3, 5.5)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('交通費合計：{{ kotsuhi_gokei }}円　　補助申請額：{{ hoshu_gaku }}円')
    run.font.size = Pt(11)
    run.font.bold = True

    path = os.path.join(ROOT, 'yytan_kotsuhi_template.docx')
    doc.save(path)
    print('ok: yytan_kotsuhi_template.docx')


if __name__ == '__main__':
    import sys
    sys.stdout.reconfigure(encoding='utf-8')
    create_jutaku_konyuu()
    create_telework_ijuu()
    create_yytan_kotsuhi()
    print('全テンプレート生成完了')

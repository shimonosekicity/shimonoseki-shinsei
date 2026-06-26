# -*- coding: utf-8 -*-
"""移住支援関連のdocxテンプレートを生成するスクリプト。
リポジトリルートで `python tools/create_ijuu_templates.py` を実行する。
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

ROOT = os.path.join(os.path.dirname(__file__), "..")


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_table_border(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def label_cell(cell, text):
    set_cell_bg(cell, 'EBF0F8')
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.bold = True
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def value_cell(cell, placeholder):
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(placeholder)
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)
    return p


def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        cells = row.cells
        if col_idx < len(cells):
            cells[col_idx].width = Cm(width_cm)


# ──────────────────────────────────────────────────────────────
# 1. 暮らしサポート補助金交付申請書
# ──────────────────────────────────────────────────────────────
def create_kurashi_support():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)

    # タイトル
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('様式第１号（第４条関係）')
    run.font.size = Pt(9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('{{ shinsei_date }}')
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run('下関市暮らしサポート補助金　交付申請書')
    run.font.size = Pt(16)
    run.font.bold = True

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('下関市長　様')
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(
        '　下関市暮らしサポート補助金交付要綱第４条の規定により、'
        '下記のとおり補助金の交付を申請します。'
    )
    run.font.size = Pt(10)

    # 申請者情報テーブル
    add_section_heading(doc, '■ 申請者情報')
    tbl = doc.add_table(rows=3, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(tbl)

    data = [
        ('氏　　名', '{{ shimei }}', 'ふりがな', '{{ shimei_kana }}'),
        ('住　　所', '{{ jusho }}', '電話番号', '{{ denwa }}'),
        ('転入年月日', '{{ tennyu_date }}', '', ''),
    ]
    for i, (l1, v1, l2, v2) in enumerate(data):
        row = tbl.rows[i]
        label_cell(row.cells[0], l1)
        value_cell(row.cells[1], v1)
        label_cell(row.cells[2], l2)
        value_cell(row.cells[3], v2)

    set_col_width(tbl, 0, 3.0)
    set_col_width(tbl, 1, 5.5)
    set_col_width(tbl, 2, 3.0)
    set_col_width(tbl, 3, 5.5)

    doc.add_paragraph()

    # 住居・家賃情報テーブル
    add_section_heading(doc, '■ 住居・家賃情報')
    tbl2 = doc.add_table(rows=2, cols=4)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(tbl2)

    data2 = [
        ('入居年月日', '{{ nyukyo_date }}', '家賃月額（円）', '{{ yachin_getsugaku }}'),
        ('補助対象月数', '{{ hoshu_tsuki }}ヶ月', '家賃支払月の開始月', '{{ yachin_start_month }}'),
    ]
    for i, (l1, v1, l2, v2) in enumerate(data2):
        row = tbl2.rows[i]
        label_cell(row.cells[0], l1)
        value_cell(row.cells[1], v1)
        label_cell(row.cells[2], l2)
        value_cell(row.cells[3], v2)

    set_col_width(tbl2, 0, 3.0)
    set_col_width(tbl2, 1, 5.5)
    set_col_width(tbl2, 2, 3.5)
    set_col_width(tbl2, 3, 5.0)

    path = os.path.join(ROOT, 'kurashi_support_template.docx')
    doc.save(path)
    print('ok: kurashi_support_template.docx')


# ──────────────────────────────────────────────────────────────
# 2. 移住支援金支給申請書
# ──────────────────────────────────────────────────────────────
def create_ijuu_shien():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('様式第１号（第６条関係）')
    run.font.size = Pt(9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('{{ shinsei_date }}')
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run('下関市移住支援金　支給申請書')
    run.font.size = Pt(16)
    run.font.bold = True

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('下関市長　様')
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(
        '　下関市移住支援事業実施要綱第６条の規定により、'
        '下記のとおり移住支援金の支給を申請します。'
    )
    run.font.size = Pt(10)

    # 申請者情報
    add_section_heading(doc, '■ 申請者情報')
    tbl = doc.add_table(rows=5, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(tbl)

    data = [
        ('氏　　名', '{{ shimei }}', 'ふりがな', '{{ shimei_kana }}'),
        ('住　　所', '{{ jusho }}', '電話番号', '{{ denwa }}'),
        ('生年月日', '{{ seinengappi }}', '年　　齢', '{{ nenrei }}歳'),
        ('転入年月日', '{{ tennyu_date }}', '転入前住所\n（東京圏の住所）', '{{ mae_jusho }}'),
        ('移住前の\n在住期間', '{{ zaiju_kikan }}', '職業（転入前）', '{{ mae_shokugyo }}'),
    ]
    for i, (l1, v1, l2, v2) in enumerate(data):
        row = tbl.rows[i]
        label_cell(row.cells[0], l1)
        value_cell(row.cells[1], v1)
        label_cell(row.cells[2], l2)
        value_cell(row.cells[3], v2)

    set_col_width(tbl, 0, 3.0)
    set_col_width(tbl, 1, 5.5)
    set_col_width(tbl, 2, 3.0)
    set_col_width(tbl, 3, 5.5)

    doc.add_paragraph()

    # 申請区分
    add_section_heading(doc, '■ 申請区分・就業情報')
    tbl2 = doc.add_table(rows=4, cols=4)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(tbl2)

    data2 = [
        ('申請区分', '{{ shinsei_kubun }}', '帯同する\n18歳未満の子の数', '{{ kodomo_count }}人'),
        ('勤務先・就業先名', '{{ kinmusaki_name }}', '勤務先所在地', '{{ kinmusaki_jusho }}'),
        ('雇用形態', '{{ koyo_keitai }}', '就業年月日', '{{ shugyou_date }}'),
        ('創業・事業内容\n（創業の場合）', '{{ jigyo_naiyou }}', '', ''),
    ]
    for i, (l1, v1, l2, v2) in enumerate(data2):
        row = tbl2.rows[i]
        label_cell(row.cells[0], l1)
        value_cell(row.cells[1], v1)
        label_cell(row.cells[2], l2)
        value_cell(row.cells[3], v2)

    set_col_width(tbl2, 0, 3.0)
    set_col_width(tbl2, 1, 5.5)
    set_col_width(tbl2, 2, 3.0)
    set_col_width(tbl2, 3, 5.5)

    path = os.path.join(ROOT, 'ijuu_shien_template.docx')
    doc.save(path)
    print('ok: ijuu_shien_template.docx')


# ──────────────────────────────────────────────────────────────
# 3. 結婚新生活支援補助金交付申請書
# ──────────────────────────────────────────────────────────────
def create_kekkon_shinseikatsu():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('様式第１号（第４条関係）')
    run.font.size = Pt(9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('{{ shinsei_date }}')
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run('下関市結婚新生活支援補助金　交付申請書')
    run.font.size = Pt(16)
    run.font.bold = True

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('下関市長　様')
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(
        '　下関市結婚新生活支援補助金交付要綱第４条の規定により、'
        '下記のとおり補助金の交付を申請します。'
    )
    run.font.size = Pt(10)

    # 夫婦情報
    add_section_heading(doc, '■ 申請者（夫婦）情報')
    tbl = doc.add_table(rows=5, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(tbl)

    # ヘッダー行
    row = tbl.rows[0]
    label_cell(row.cells[0], '項　目')
    label_cell(row.cells[1], '夫')
    label_cell(row.cells[2], '妻')
    label_cell(row.cells[3], '備　考')

    data = [
        ('氏　　名', '{{ otto_shimei }}', '{{ tsuma_shimei }}', ''),
        ('ふりがな', '{{ otto_kana }}', '{{ tsuma_kana }}', ''),
        ('生年月日', '{{ otto_seinengappi }}', '{{ tsuma_seinengappi }}', ''),
        ('婚姻届受理日', '{{ konin_date }}', '', ''),
    ]
    for i, (l, v1, v2, v3) in enumerate(data):
        row = tbl.rows[i + 1]
        label_cell(row.cells[0], l)
        value_cell(row.cells[1], v1)
        value_cell(row.cells[2], v2)
        value_cell(row.cells[3], v3)

    set_col_width(tbl, 0, 3.0)
    set_col_width(tbl, 1, 4.5)
    set_col_width(tbl, 2, 4.5)
    set_col_width(tbl, 3, 5.0)

    doc.add_paragraph()

    # 新居情報
    add_section_heading(doc, '■ 新居情報')
    tbl2 = doc.add_table(rows=3, cols=4)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(tbl2)

    data2 = [
        ('新居住所', '{{ shinkyo_jusho }}', '入居年月日', '{{ nyukyo_date }}'),
        ('住居形態', '{{ jutaku_keitai }}', '世帯合計所得', '{{ setai_shotoku }}万円'),
        ('転入前住所', '{{ mae_jusho }}', '', ''),
    ]
    for i, (l1, v1, l2, v2) in enumerate(data2):
        row = tbl2.rows[i]
        label_cell(row.cells[0], l1)
        value_cell(row.cells[1], v1)
        label_cell(row.cells[2], l2)
        value_cell(row.cells[3], v2)

    set_col_width(tbl2, 0, 3.0)
    set_col_width(tbl2, 1, 5.5)
    set_col_width(tbl2, 2, 3.0)
    set_col_width(tbl2, 3, 5.5)

    doc.add_paragraph()

    # 補助申請内容
    add_section_heading(doc, '■ 補助申請内容（該当するものを記入）')
    tbl3 = doc.add_table(rows=5, cols=3)
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(tbl3)

    # ヘッダー行
    row = tbl3.rows[0]
    label_cell(row.cells[0], '費用区分')
    label_cell(row.cells[1], '支出額（円）')
    label_cell(row.cells[2], '備　考')

    cost_items = [
        ('住居賃借費（家賃・敷金・礼金等）', '{{ yachin_gaku }}', '{{ yachin_biko }}'),
        ('住居取得費（購入・建築費）', '{{ konya_gaku }}', '{{ konya_biko }}'),
        ('引越費用', '{{ hikkoshi_gaku }}', '{{ hikkoshi_biko }}'),
        ('リフォーム費用', '{{ reform_gaku }}', '{{ reform_biko }}'),
    ]
    for i, (l, v1, v2) in enumerate(cost_items):
        row = tbl3.rows[i + 1]
        label_cell(row.cells[0], l)
        value_cell(row.cells[1], v1)
        value_cell(row.cells[2], v2)

    set_col_width(tbl3, 0, 5.0)
    set_col_width(tbl3, 1, 4.0)
    set_col_width(tbl3, 2, 8.0)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('補助申請額合計：{{ hoshu_gokei }}円')
    run.font.size = Pt(11)
    run.font.bold = True

    path = os.path.join(ROOT, 'kekkon_shinseikatsu_template.docx')
    doc.save(path)
    print('ok: kekkon_shinseikatsu_template.docx')


if __name__ == '__main__':
    import sys
    sys.stdout.reconfigure(encoding='utf-8')
    create_kurashi_support()
    create_ijuu_shien()
    create_kekkon_shinseikatsu()
    print('全テンプレート生成完了')

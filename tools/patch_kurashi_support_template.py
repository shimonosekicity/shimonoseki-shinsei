# -*- coding: utf-8 -*-
"""公式Word様式にJinja2プレースホルダーを埋め込むスクリプト。
リポジトリルートで `python tools/patch_kurashi_support_template.py` を実行する。
"""
import os, copy
from docx import Document
from docx.oxml.ns import qn

ROOT = os.path.join(os.path.dirname(__file__), "..")
SRC  = os.path.join(ROOT, "kurashi_support_template.docx")
DEST = os.path.join(ROOT, "kurashi_support_template.docx")


def set_run_text(run, text):
    run.text = text


def patch():
    doc = Document(SRC)

    # ── 段落 ──────────────────────────────────────────────
    # [1] 申請日
    p1 = doc.paragraphs[1]
    p1.runs[0].text = "{{ shinsei_date }}"

    # [5] 住所
    p5 = doc.paragraphs[5]
    p5.runs[1].text = "住所　{{ jusho }}"

    # [6] 氏名
    p6 = doc.paragraphs[6]
    p6.runs[1].text = "氏名　{{ shimei }}（{{ shimei_kana }}）　"

    # [7] 電話
    p7 = doc.paragraphs[7]
    p7.runs[1].text = "連絡先（電話）　{{ denwa }}"

    # ── テーブル ──────────────────────────────────────────
    tbl = doc.tables[0]

    # 行0: 賃貸借契約日
    tbl.rows[0].cells[1].paragraphs[0].runs[0].text = "{{ chinshaku_date }}"

    # 行1: 入居年月日
    tbl.rows[1].cells[1].paragraphs[0].runs[0].text = "{{ nyukyo_date }}"

    # 行2: 家賃の月額
    tbl.rows[2].cells[1].paragraphs[1].runs[0].text = "月額　{{ yachin_getsugaku }}　円"

    # 行3: 補助金申請額 — ①②③④ を書き換え
    cell3 = tbl.rows[3].cells[1]
    paras = cell3.paragraphs

    # ① (p[1])
    for r in paras[1].runs:
        r.text = ""
    paras[1].runs[0].text = "①　家賃の月額×１／２　　{{ yachin_half }}　円"

    # ② (p[3])
    for r in paras[3].runs:
        r.text = ""
    paras[3].runs[0].text = "②　補助金額（月額）　{{ hoshu_gaku_monthly }}　円"

    # ③ (p[8])
    for r in paras[8].runs:
        r.text = ""
    paras[8].runs[0].text = (
        "③　家賃の支払月　{{ yachin_start_month }}　～　{{ yachin_end_month }}"
    )

    # ⇒ ヶ月 (p[9])
    for r in paras[9].runs:
        r.text = ""
    paras[9].runs[0].text = "⇒　{{ hoshu_tsuki }}ヶ月"

    # ④ (p[11])
    for r in paras[11].runs:
        r.text = ""
    paras[11].runs[0].text = (
        "④　補助金申請額（②×③）　{{ hoshu_gaku_monthly }}円　×　{{ hoshu_tsuki }}ヶ月"
    )

    # 補助金申請額合計 (p[13])
    for r in paras[13].runs:
        r.text = ""
    paras[13].runs[0].text = "補助金申請額　{{ hoshu_gokei }}　円"

    doc.save(DEST)
    print("ok: kurashi_support_template.docx")


if __name__ == "__main__":
    import sys
    sys.stdout.reconfigure(encoding="utf-8")
    patch()

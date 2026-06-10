# -*- coding: utf-8 -*-
"""docxテンプレートに不足しているプレースホルダを一括追加するスクリプト。

これまで「ダウンロード後にWordへ直接記入」だった項目を差し込み可能にする。
リポジトリルートで `python tools/fix_templates.py` を実行する（再実行しても安全）。
"""
import copy
import os
import re
import sys

from docx import Document
from docx.oxml.ns import qn

ROOT = os.path.join(os.path.dirname(__file__), "..")


def find_para(doc, pred, label):
    for i, p in enumerate(doc.paragraphs):
        if pred(p.text):
            return i, p
    raise SystemExit(f"NOT FOUND: {label}")


def eq(text):
    return lambda t: t.strip() == text


def starts(text):
    return lambda t: t.strip().startswith(text)


def append_run(p, text):
    if text in p.text:
        return  # 既に適用済み
    run = p.add_run(text)
    prev = [r for r in p.runs[:-1] if r.text.strip()]
    if prev and prev[-1].element.rPr is not None:
        run.element.insert(0, copy.deepcopy(prev[-1].element.rPr))


def fill_empty_after(doc, idx, text):
    """idx番の段落より後ろで最初の空段落にtextを入れる。"""
    for p in doc.paragraphs[idx + 1:]:
        if text in p.text:
            return
        if not p.text.strip():
            p.add_run(text)
            return
    raise SystemExit(f"no empty paragraph after #{idx} for {text}")


def rebuild_para(p, new_text):
    """段落全体のテキストを置き換える（先頭ランの書式を維持）。"""
    if new_text == p.text:
        return
    runs = p.runs
    if not runs:
        p.add_run(new_text)
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""


def replace_in_cell(cell, old, new, count=1):
    done = 0
    for p in cell.paragraphs:
        if new in p.text:
            return  # 適用済み
        for r in p.runs:
            while old in r.text and done < count:
                r.text = r.text.replace(old, new, 1)
                done += 1
        if done >= count:
            return
    if done == 0:
        raise SystemExit(f"checkbox not found in cell: {cell.text!r}")


def set_cell_text(cell, text):
    if text in cell.text:
        return
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)
    for extra in cell.paragraphs[1:]:
        for r in extra.runs:
            r.text = ""


def unique_cells(row):
    seen, out = set(), []
    for c in row.cells:
        if id(c._tc) not in seen:
            seen.add(id(c._tc))
            out.append(c)
    return out


CK_KEYS = ["gozen", "encho1", "gogo", "encho2", "yakan"]


def patch_schedule_table(table, row_start, row_end):
    """公民館の使用日程表（5行）に日付・曜日・室名・時間帯チェック・摘要を差し込む。"""
    for n, ri in enumerate(range(row_start, row_end + 1), start=1):
        cells = unique_cells(table.rows[ri])
        assert len(cells) == 10, f"expected 10 cells, got {len(cells)}: {[c.text for c in cells]}"
        num, date, youbi, room, *checks, biko = cells
        set_cell_text(date, f"{{{{ date_{n} }}}}")
        # 元様式の箇条書き記号「・」が日付の前に出ないように numPr を除去
        for p in date.paragraphs:
            pPr = p._p.pPr
            if pPr is not None:
                numPr = pPr.find(qn("w:numPr"))
                if numPr is not None:
                    pPr.remove(numPr)
        set_cell_text(youbi, f"{{{{ youbi_{n} }}}}")
        set_cell_text(room, f"{{{{ room_{n} }}}}")
        for key, cell in zip(CK_KEYS, checks):
            replace_in_cell(cell, "□", f"{{{{ ck_{n}_{key} }}}}")
        set_cell_text(biko, f"{{{{ biko_{n} }}}}")


def patch_kouminkan_shiyo():
    path = os.path.join(ROOT, "kouminkan_shiyo_template.docx")
    doc = Document(path)
    patch_schedule_table(doc.tables[3], 4, 8)
    # 社会教育法第23条の確認チェック（表の右端の□）
    confirm_cell = unique_cells(doc.tables[3].rows[0])[-1]
    set_cell_text(confirm_cell, "{{ houki_check }}")
    doc.save(path)
    print("ok: kouminkan_shiyo")


def patch_kouminkan_genmen():
    path = os.path.join(ROOT, "kouminkan_genmen_template.docx")
    doc = Document(path)
    # 申請者表の「団体名」欄が空
    set_cell_text(unique_cells(doc.tables[0].rows[1])[2], "{{ dantai_name }}")
    # 使用目的欄が空
    cells = unique_cells(doc.tables[1].rows[1])
    target = next(c for c in cells if not c.text.strip() or "shiyou_mokuteki" in c.text)
    set_cell_text(target, "{{ shiyou_mokuteki }}")
    patch_schedule_table(doc.tables[1], 4, 8)
    doc.save(path)
    print("ok: kouminkan_genmen")


def patch_kouminkan_chushi():
    path = os.path.join(ROOT, "kouminkan_chushi_template.docx")
    doc = Document(path)
    patch_schedule_table(doc.tables[1], 2, 6)
    doc.save(path)
    print("ok: kouminkan_chushi")


def patch_suido():
    path = os.path.join(ROOT, "suido_syoyusha_template.docx")
    doc = Document(path)
    _, p = find_para(doc, starts("氏　名"), "suido 氏名")
    append_run(p, "　{{ shutsu_shimei }}")
    doc.save(path)
    print("ok: suido")


def patch_shogai():
    path = os.path.join(ROOT, "shogai_jigyo_template.docx")
    doc = Document(path)
    # 「次のとおり　　…　を開始します」の空白部分（1ページ目のみ）
    _, p = find_para(doc, lambda t: t.startswith("次のとおり"), "shogai 次のとおり")
    if "jigyo_shurui" not in p.text:
        rebuild_para(p, re.sub("　{4,}", "　{{ jigyo_shurui }}　", p.text, count=1))
    appends = [
        ("（１）種類", "jigyo_shurui"),
        ("（２）内容", "jigyo_naiyou"),
        ("（１）氏名（名　称）", "keieisha_shimei"),
        ("（２）住所（所在地）", "keieisha_jusho"),
        ("当該市町の名称を含む。）", "jigyo_kuiki"),
        ("(１)　名　　　称", "shisetsu_meisho"),
        ("(２)　種　　　類", "shisetsu_shurui"),
        ("(３)　所　在　地", "shisetsu_shozaichi"),
        ("(４)　入所定員", "nyusho_teiin"),
    ]
    for text, var in appends:
        _, p = find_para(doc, starts(text), f"shogai {text}")
        append_run(p, f"　{{{{ {var} }}}}")
    doc.save(path)
    print("ok: shogai")


def patch_noshin():
    path = os.path.join(ROOT, "noshin_moushide_template.docx")
    doc = Document(path)
    fills = [
        (eq("理由（変更理由及び早急に変更をする理由等）"), "{{ henko_riyu }}"),
        (eq("３．当該地を選定した理由"), "{{ sentei_riyu }}"),
        (lambda t: t.strip().startswith("ェ．その他（"), "{{ nochi_jokyo }}"),
        (eq("ィ．　　　無"), "{{ koko_toshi }}"),
        (eq("汚悪水の処理方法"), "{{ osui_shori }}"),
        (eq("日照関係"), "{{ nicchou }}"),
        (eq("その他"), "{{ sonota_eikyou }}"),
    ]
    for pred, text in fills:
        idx, _ = find_para(doc, pred, text)
        fill_empty_after(doc, idx, text)
    appends = [
        (starts("氏名："), "{{ shutai_shimei }}"),
        (starts("住所："), "{{ shutai_jusho }}"),
        (starts("地権者との関係："), "{{ shutai_kankei }}"),
        (starts("（２）専業、兼業の区別"), "　{{ senkyo_kenbyo }}"),
    ]
    for pred, text in appends:
        _, p = find_para(doc, pred, text)
        append_run(p, text)
    # 「７　事業主体（　…　）の農業経営の状況」の括弧内
    _, p = find_para(doc, lambda t: t.strip().startswith("７．事業主体（"), "noshin ７事業主体")
    if "shutai_shimei" not in p.text:
        rebuild_para(p, re.sub("（　*）", "（{{ shutai_shimei }}）", p.text, count=1))
    # 「（総事業費　…　円）」
    _, p = find_para(doc, lambda t: "（総事業費" in t, "noshin 総事業費")
    if "jigyo_hi" not in p.text:
        rebuild_para(p, re.sub("総事業費　*円", "総事業費　{{ jigyo_hi }}　円", p.text, count=1))
    # 表：用地面積（区域内・区域外・計）
    cells = unique_cells(doc.tables[1].rows[2])
    set_cell_text(cells[0], "{{ kuiki_nai_menseki }}")
    set_cell_text(cells[1], "{{ kuiki_gai_menseki }}")
    set_cell_text(cells[-1], "{{ kuiki_kei }}")
    # 表：農業経営面積（田・畑の現在／変更後／増減）
    tbl = doc.tables[2]
    for ri, suffix in [(1, "genzai"), (2, "ato"), (3, "zougen")]:
        cells = unique_cells(tbl.rows[ri])
        set_cell_text(cells[1], f"{{{{ ta_{suffix} }}}}")
        set_cell_text(cells[2], f"{{{{ hata_{suffix} }}}}")
    doc.save(path)
    print("ok: noshin")


if __name__ == "__main__":
    sys.stdout.reconfigure(encoding="utf-8")
    patch_kouminkan_shiyo()
    patch_kouminkan_genmen()
    patch_kouminkan_chushi()
    patch_suido()
    patch_shogai()
    patch_noshin()
